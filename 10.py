import requests
import urllib3
from docx import Document
from docx.shared import Pt

urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

# 设置 API 地址和组 ID
api_url = 'https://172.21.5.27:3443/api/v1/scans'
group_id = '7b6d6dd0-3a99-406f-83c0-49f003ad3e89'

# 设置 API 认证信息
api_key = "1986ad8c0a5b3df4d7028d5f3c06e936ca723deb5d73441da9d991805c93bfd00"
headers = {"X-Auth": api_key}

# 发送 GET 请求获取任务信息
response = requests.get(f'{api_url}?group_id={group_id}', headers=headers, verify=False)

# 获取返回的 JSON 数据并提取 scan_id 和 scan_session_id
scan_data = response.json()['scans']
scan_id_list = []
scan_session_id_list = []
for scan in scan_data:
    # 判断 IP 是否以 10.252 开头
    if str(scan['target']).startswith('10.252.'):
        scan_id_list.append(scan['scan_id'])
        scan_session_id_list.append(scan['current_session']['scan_session_id'])

# 创建 Word 文档对象
document = Document()

# 定义数据列表
data_list = []
for scan_id, scan_session_id in zip(scan_id_list, scan_session_id_list):
    vulnerabilities_url = f'{api_url}/{scan_id}/results/{scan_session_id}/vulnerabilities'
    response = requests.get(vulnerabilities_url, headers=headers, verify=False)
    vulnerabilities_data = response.json()
    for vuln in vulnerabilities_data['vulnerabilities']:
        vuln_id = vuln['vuln_id']
        vuln_url = f'{api_url}/{scan_id}/results/{scan_session_id}/vulnerabilities/{vuln_id}'
        response = requests.get(vuln_url, headers=headers, verify=False)
        vuln_details = response.json()
        # 对 cvss_score 值进行判断
        if vuln_details['cvss_score'] >= 7.0:
            cvss_score_str = '高危'
        elif vuln_details['cvss_score'] >= 4.0:
            cvss_score_str = '中危'
        elif vuln_details['cvss_score'] >= 0.1:
            cvss_score_str = '低危'
        else:
            cvss_score_str = '信息'
        data_list.append({
            'vt_name': vuln_details['vt_name'],
            'long_description': vuln_details['long_description'],
            'affects_detail': vuln_details['affects_detail'],
            'description': vuln_details['description'],
            'impact': vuln_details['impact'],
            'affects_url': vuln_details['affects_url'],
            'request': vuln_details['request'],
            'cvss_score': cvss_score_str,
            'recommendation': vuln_details['recommendation'],
            'references': vuln_details['references'],
        })

# 添加标题
document.add_heading('漏洞列表', 0)

# 遍历数据列表，输出数据到文档中
for index, data in enumerate(data_list):
    # 生成段落
    paragraph = document.add_paragraph()
    # 设置段落的格式和样式
    paragraph.space_after = Pt(10)  # 段后距离
    paragraph.space_before = Pt(10)  # 段前距离
    paragraph.font.name = '宋体' # 设置段落中的字体
    paragraph.font.size = Pt(12) # 设置段落中的字号
    # 合并段落中的各个字段
    paragraph.add_run(f"漏洞名称：{data['vt_name']}\n漏洞描述：{data['long_description']}\n" \
                           f"影响版本：{data['affects_detail']}\n漏洞详情：{data['description']}\n" \
                           f"风险评估：{data['impact']}\n相关 URL：{data['affects_url']}\n" \
                           f"请求数据：{data['request']}\n" \
                           f"CVSS评分：{data['cvss_score']}\n建议处理：{data['recommendation']}\n参考链接：{data['references']}\n")
    # 添加换行符
    paragraph.add_run('\n')

# 保存文档
document.save('D:\\pythonProject7\\output57.docx')
